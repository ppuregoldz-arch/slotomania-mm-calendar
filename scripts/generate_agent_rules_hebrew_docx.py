#!/usr/bin/env python3
"""Generate English Word doc — MM Calendar promo guidelines (field template + categories)."""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm

from promo_guidelines_doc_data import CATEGORY_ORDER, GUIDELINES
from promo_guidelines_doc_render import add_para, add_promo_block

OUT = (
    Path(__file__).resolve().parents[1]
    / "mm_calendar"
    / "documentation"
    / "sm_agent_calendar_prediction_rules_en.docx"
)
DESKTOP_COPY = Path.home() / "Desktop" / "sm_agent_calendar_prediction_rules_en.docx"
LEGACY_HE_COPY = (
    Path(__file__).resolve().parents[1]
    / "mm_calendar"
    / "documentation"
    / "sm_agent_calendar_prediction_rules_he.docx"
)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def build() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    title = doc.add_heading("MM Calendar — Promo Guidelines (field template)", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(
        doc,
        f"Generated {date.today().strftime('%Y-%m-%d')} · Sources: mm_calendar/, real_months.json, "
        "August 2026 Monday board. On conflict, monthly_guidelines/YYYY-MM.md wins.",
    )

    add_heading(doc, "Short introduction", 1)
    for item in [
        "Promos are grouped by lane: Purchase, Core, Gems, Dice, Clan/ADS, Seasonal, Events, Other.",
        "What promotions do we have = five different named SKUs/variants (never the same title five times).",
        "Examples on calendar = five different dates/instances from Monday or plan history.",
        "Global timing: Promo Time 11:00 UTC · Night Plan 00:00 UTC · Time-limited ≥12:00 UTC.",
        "Real second VFM (not Clan-Dash): RYD / Rolling / Buy All / Decoy / Prize Mania (+ Coin Sale on sale days).",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    by_cat: dict[str, list] = {c: [] for c in CATEGORY_ORDER}
    for entry in GUIDELINES:
        by_cat.setdefault(entry["category"], []).append(entry)

    add_heading(doc, "Guidelines by category", 1)
    for cat in CATEGORY_ORDER:
        items = by_cat.get(cat) or []
        if not items:
            continue
        add_heading(doc, cat, 2)
        for entry in items:
            add_promo_block(doc, entry, heading_level=3)

    add_heading(doc, "Appendix — authority & files", 1)
    add_para(
        doc,
        "Priority: monthly_guidelines → constraints + rules_cheatsheet → learnings → prediction/calibrate.",
        bold=True,
    )
    for item in [
        "scripts/promo_guidelines_doc_data.py · scripts/promo_guidelines_catalog.py",
        "scripts/generate_agent_rules_hebrew_docx.py · scripts/generate_promo_subject_docx.py",
        "scripts/build_august_2026_plan.py · scripts/upload_mm_calendar_day_monday.py",
        "AGENTS.md · TEAM_WORKLOG.md after sessions that change files",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.save(str(OUT))
    DESKTOP_COPY.write_bytes(OUT.read_bytes())
    LEGACY_HE_COPY.write_bytes(OUT.read_bytes())
    print(f"Wrote {OUT}")
    print(f"Copied to {DESKTOP_COPY}")


if __name__ == "__main__":
    build()
