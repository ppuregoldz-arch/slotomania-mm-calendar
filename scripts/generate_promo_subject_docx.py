#!/usr/bin/env python3
"""Generate manager fill-in Word files per MM Calendar subject (Hebrew UX + AI_FIELD tags)."""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt

from promo_guidelines_doc_data import GUIDELINES, PromoGuideline
from promo_guidelines_doc_render import add_manager_cover, add_manager_promo_block

DOC_ROOT = Path(__file__).resolve().parents[1] / "mm_calendar" / "documentation" / "subjects"
DESKTOP_ROOT = Path.home() / "Desktop" / "mm_calendar_subject_docs"

# (English key, filename, Hebrew title, owner hint, promo name keys…)
SUBJECTS: tuple[tuple[str, str, str, str, tuple[str, ...]], ...] = (
    (
        "Core",
        "sm_mm_calendar_core_MANAGER.docx",
        "ליבה — אתגרי Core ו-Custom Pod",
        "מנהל/ת Core / Gameplay",
        (
            "Core challenges (Win Master / PYP / Spin Zone / Ace Heist / MES)",
            "Custom Pod",
        ),
    ),
    (
        "Album",
        "sm_mm_calendar_album_MANAGER.docx",
        "אלבום — Shiny Show וקצב אלבום",
        "מנהל/ת Album / Shiny",
        (
            "Shiny Show",
            "Album rhythm",
        ),
    ),
    (
        "Purchase Tools — Offers + Daily Deals",
        "sm_mm_calendar_purchase_offers_daily_deals_MANAGER.docx",
        "כלי רכישה — הצעות + Daily Deals",
        "מנהל/ת Purchase Offers / DD",
        (
            "Daily Deal (DD)",
            "RYD (Reveal Your Deal)",
            "Buy All / Mystery Buy All",
            "Decoy Bonanza",
            "Rolling Offer — Buy X Get Y (BXGY)",
            "Rolling Offer — Buy More for Less (BMFL)",
            "Coin Sale",
            "MGAP (BOGO / Bigger / Matched / Wild)",
            "Prize Mania",
            "Counter PO",
            "Fortune Dip / Extreme Stamp",
            "Price Cut 20%",
            "SlotoBucks / 1st of month denom",
            "Boosted Gemback",
            "x2 GGS",
            "Gems Sale",
            "Popup Store (test)",
        ),
    ),
    (
        "Purchase Tools — Features",
        "sm_mm_calendar_purchase_features_MANAGER.docx",
        "כלי רכישה — פיצ'רים (Golden Spin, Piggy, Dice)",
        "מנהל/ת Purchase Features",
        (
            "Golden Spin",
            "Piggy",
            "Dice Deluxe / Dice Booster",
        ),
    ),
    (
        "Happy Hours",
        "sm_mm_calendar_happy_hours_MANAGER.docx",
        "Happy Hours — HH, Lotto, אירועים",
        "מנהל/ת Events / HH / Lotto",
        (
            "Happy Hour / Jumbo Giveaway",
            "Lotto peak + LBP",
            "Black Diamond / Event anchors",
        ),
    ),
    (
        "Mid Term",
        "sm_mm_calendar_mid_term_MANAGER.docx",
        "Mid Term — Winovate, Mega Pods, רוטטורים",
        "מנהל/ת Mid Term",
        ("Mid Term (Winovate / Mega Pods / Quest·Globez·Figz)",),
    ),
    (
        "Short Term",
        "sm_mm_calendar_short_term_MANAGER.docx",
        "Short Term — Blast / Battlesheep / SNL",
        "מנהל/ת Short Term / עונות",
        ("Short Term (Blast / Battlesheep / SNL)",),
    ),
    (
        "Clan & Dash",
        "sm_mm_calendar_clan_dash_MANAGER.docx",
        "Clan & Dash + ADS",
        "מנהל/ת Clan-Dash / ADS",
        (
            "Clan-Dash / Dash Pass",
            "ADS PO",
        ),
    ),
)


def _index_guidelines() -> dict[str, PromoGuideline]:
    return {g["name"]: g for g in GUIDELINES}


def build_subject_doc(
    title_en: str,
    title_he: str,
    owner_hint: str,
    entries: list[PromoGuideline],
) -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    add_manager_cover(
        doc,
        subject_title_en=title_en,
        subject_title_he=title_he,
        owner_hint=owner_hint,
        promo_count=len(entries),
    )

    doc.add_heading(f"פרומואים בתחום — {title_he}", 1)
    add_para = doc.add_paragraph
    p = add_para()
    run = p.add_run(
        f"אחראי מוצע: {owner_hint} · "
        "מלאו כל בלוק; שמרו את הקובץ בשם המקורי עם _MANAGER."
    )
    run.font.size = Pt(11)

    if not entries:
        doc.add_paragraph("(אין מיפוי — עדכנו scripts/promo_guidelines_doc_data.py.)")
    for g in entries:
        add_manager_promo_block(doc, g)
    return doc


def main() -> None:
    by_name = _index_guidelines()
    DOC_ROOT.mkdir(parents=True, exist_ok=True)
    DESKTOP_ROOT.mkdir(parents=True, exist_ok=True)

    written: list[str] = []
    for title_en, filename, title_he, owner_hint, names in SUBJECTS:
        entries = [by_name[n] for n in names if n in by_name]
        missing = [n for n in names if n not in by_name]
        if missing:
            print(f"WARN {title_en}: missing keys {missing}")
        doc = build_subject_doc(title_en, title_he, owner_hint, entries)
        out_repo = DOC_ROOT / filename
        out_desktop = DESKTOP_ROOT / filename
        doc.save(str(out_repo))
        out_desktop.write_bytes(out_repo.read_bytes())
        written.append(f"{title_he} → {filename}")

    print(f"Manager forms ({date.today().isoformat()}):")
    print(f"  {DOC_ROOT}")
    print(f"  {DESKTOP_ROOT}")
    for line in written:
        print(f"  · {line}")


if __name__ == "__main__":
    main()
