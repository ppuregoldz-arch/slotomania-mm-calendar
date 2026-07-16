"""Shared Word rendering — agent reference blocks + manager fill-in forms."""

from __future__ import annotations

import re
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from promo_guidelines_doc_data import PromoGuideline

_FILL_LINE = "______________________________________________________________"
_FILL_SHORT = "______________________________"


def promo_slug(name: str) -> str:
    s = name.lower()
    s = re.sub(r"[^a-z0-9]+", "_", s)
    return s.strip("_")[:72] or "promo"


def add_para(doc: Document, text: str, bold: bool = False, *, size: int = 11, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)


def add_ai_field_tag(doc: Document, field_key: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(f"AI_FIELD: {field_key}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_fill_space(doc: Document, lines: int = 2) -> None:
    for _ in range(lines):
        add_para(doc, _FILL_LINE, size=10)


def add_field(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label} ")
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(value)
    r2.font.size = Pt(11)


def _unique_lines(lines: Iterable[str], limit: int = 5) -> list[str]:
    seen: set[str] = set()
    out: list[str] = []
    for line in lines:
        key = line.strip()
        if not key or key in seen:
            continue
        seen.add(key)
        out.append(key)
        if len(out) >= limit:
            break
    return out


def add_numbered_lines(doc: Document, lines: list[str], *, max_items: int = 5) -> None:
    uniq = _unique_lines(lines, max_items)
    for i, line in enumerate(uniq, 1):
        doc.add_paragraph(f"{i}. {line}", style="List Number")
    for j in range(len(uniq) + 1, max_items + 1):
        doc.add_paragraph(f"{j}. (Add distinct variant — do not duplicate row {j - 1})", style="List Number")


# --- Manager-facing form (subject Word files for monetization owners) ---


def add_manager_cover(
    doc: Document,
    *,
    subject_title_en: str,
    subject_title_he: str,
    owner_hint: str,
    promo_count: int,
) -> None:
    h = doc.add_heading("טופס הנחיות מוניטיזציה — MM Calendar", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, f"תחום: {subject_title_he}", bold=True, size=14)
    add_para(doc, f"Subject: {subject_title_en}", italic=True)
    add_para(doc, f"נוצר: {__import__('datetime').date.today().isoformat()} · {promo_count} סוגי פרומואים בקובץ")

    doc.add_heading("פרטי ממלא/ת", 1)
    add_ai_field_tag(doc, "form.owner_name")
    add_para(doc, "שם מלא:", bold=True)
    add_fill_space(doc, 1)
    add_ai_field_tag(doc, "form.owner_email")
    add_para(doc, "אימייל / Teams:", bold=True)
    add_fill_space(doc, 1)
    add_ai_field_tag(doc, "form.completed_date")
    add_para(doc, "תאריך מילוי:", bold=True)
    add_para(doc, _FILL_SHORT)
    add_ai_field_tag(doc, "form.status")
    add_para(doc, "סטטוס:  ☐ טיוטה   ☐ מוכן להחזרה ל-Cursor / AI", bold=True)

    doc.add_heading("איך למלא (2 דקות)", 1)
    for item in [
        "מלאו רק בחלקים המסומנים «מילוי שלכם». אפשר למחוק רפרנס אם לא רלוונטי ולכתוב במקומו.",
        "ב«וריאנטים» — חמש שורות שונות (לא אותו פרומו 5 פעמים). כותרת כמו במאנדיי.",
        "ב«דוגמאות בלוח» — חמש תאריכים שונים + שם שורה אמיתי מההיסטוריה.",
        "ב«רפרנסים» — לינק למאנדיי, חודש לדוגמה (2026-08), Confluence, קובץ Excel.",
        "אל תשנו את שורות AI_FIELD — הן עוזרות לסוכן לקרוא את הקובץ בחזרה.",
        "החזירו את הקובץ השמור (.docx) לליד / ל-Cursor כקובץ מצורף.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("תזמון גלובלי (לכולם)", 2)
    for item in [
        "Promo Time = 11:00 UTC · Night Plan = 00:00 UTC.",
        "פרומואים מוגבלים בזמן — בדרך כלל מ-12:00 UTC.",
        "Promo Time ריק במאנדיי → השורה מוסתרת בתצוגת לוח.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_page_break()


def add_manager_fill_field(
    doc: Document,
    *,
    field_key: str,
    he_label: str,
    en_hint: str = "",
    reference: str = "",
    fill_lines: int = 2,
) -> None:
    add_para(doc, he_label, bold=True)
    if en_hint:
        add_para(doc, en_hint, italic=True, size=10)
    add_ai_field_tag(doc, field_key)
    if reference:
        add_para(doc, "רפרנס מהמערכת (אפשר לערוך או למחוק):", bold=True, size=10)
        add_para(doc, reference, italic=True, size=10)
    add_para(doc, "מילוי שלכם:", bold=True, size=10)
    add_fill_space(doc, fill_lines)


def add_manager_numbered_fill(
    doc: Document,
    *,
    field_key_prefix: str,
    he_label: str,
    reference_lines: list[str],
    count: int = 5,
) -> None:
    add_para(doc, he_label, bold=True)
    add_para(
        doc,
        "חשוב: כל שורה = פרומו/דוגמה אחרת (לא לשכפל את אותו שם).",
        italic=True,
        size=10,
    )
    refs = _unique_lines(reference_lines, count)
    if refs:
        add_para(doc, "רפרנס מהמערכת (לעיון):", bold=True, size=10)
        for i, line in enumerate(refs, 1):
            doc.add_paragraph(f"{i}. {line}", style="List Number")
    add_para(doc, "מילוי שלכם:", bold=True, size=10)
    for i in range(1, count + 1):
        add_ai_field_tag(doc, f"{field_key_prefix}_{i}")
        add_para(doc, f"{i}. ", bold=True, size=10)
        add_para(doc, _FILL_LINE, size=10)


def add_manager_promo_block(doc: Document, g: PromoGuideline) -> None:
    slug = promo_slug(g["name"])
    doc.add_heading(f"[PROMO_ID: {slug}] {g['name']}", 2)

    add_manager_fill_field(
        doc,
        field_key=f"promo.{slug}.name",
        he_label="שם סוג הפרומו (Name)",
        en_hint="Official family name as in Monday / CRM.",
        reference=g["name"],
        fill_lines=1,
    )
    add_manager_fill_field(
        doc,
        field_key=f"promo.{slug}.times_per_week",
        he_label="תדירות / ימים ספציפיים (Times per week / specific day)",
        reference=g["times"],
    )
    add_manager_fill_field(
        doc,
        field_key=f"promo.{slug}.duration",
        he_label="משך (Duration)",
        reference=g["duration"],
    )
    add_manager_fill_field(
        doc,
        field_key=f"promo.{slug}.lane",
        he_label="ליין: Core / Gems / Offer / SB / HH / Mid Term / Seasonal",
        reference=g["lane"],
        fill_lines=1,
    )
    add_manager_fill_field(
        doc,
        field_key=f"promo.{slug}.cant_go_live_with",
        he_label="אסור לצאת יחד עם (Can't go live with)",
        reference=g["cant_with"],
    )

    catalog = list(g.get("promotion_catalog") or ())
    if not catalog and g.get("variants"):
        catalog = [g["variants"]]
    add_manager_numbered_fill(
        doc,
        field_key_prefix=f"promo.{slug}.promotion_variant",
        he_label="אילו פרומואים יש לנו? (What promotions do we have — 5 וריאנטים שונים)",
        reference_lines=catalog,
    )

    add_manager_fill_field(
        doc,
        field_key=f"promo.{slug}.additional_info",
        he_label="מידע נוסף (Additional info)",
        reference=g["extra"],
        fill_lines=3,
    )

    add_manager_numbered_fill(
        doc,
        field_key_prefix=f"promo.{slug}.calendar_example",
        he_label="דוגמאות מהלוח (5 תאריכים שונים + שם שורה)",
        reference_lines=list(g.get("examples") or ()),
    )

    doc.add_heading("רפרנסים וקישורים", 3)
    add_para(doc, "לינקים, חודשים לדוגמה, מסמכים — עוזר ל-AI לאמת.", italic=True, size=10)
    for key, label in (
        (f"promo.{slug}.ref_monday", "לינק / מזהה לוח Monday:"),
        (f"promo.{slug}.ref_month", "חודש לדוגמה (YYYY-MM):"),
        (f"promo.{slug}.ref_confluence", "Confluence / מסמך / Slack:"),
        (f"promo.{slug}.ref_other", "אחר (Excel, CRM, LiveOps):"),
    ):
        add_para(doc, label, bold=True, size=10)
        add_ai_field_tag(doc, key)
        add_fill_space(doc, 1)

    add_manager_fill_field(
        doc,
        field_key=f"promo.{slug}.manager_notes",
        he_label="הערות פתוחות / פערים / שאלות לליד",
        en_hint="Anything the calendar agent should know that is not above.",
        fill_lines=4,
    )

    p = doc.add_paragraph()
    run = p.add_run(f"=== PROMO_BLOCK_END | id={slug} ===")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    doc.add_paragraph()


# --- Compact blocks for combined agent reference doc ---


def add_promo_block(doc: Document, g: PromoGuideline, *, heading_level: int = 3) -> None:
    doc.add_heading(g["name"], level=heading_level)
    add_field(doc, "Name -", g["name"])
    add_field(doc, "Times Per Week/Specific Day -", g["times"])
    add_field(doc, "Duration -", g["duration"])
    add_field(doc, "Core/Gems/Offer/SB/HH/Mid Term/Seasonal -", g["lane"])
    add_field(doc, "Can't go live with -", g["cant_with"])
    add_field(doc, "What promotions do we have -", "")
    catalog = list(g.get("promotion_catalog") or ())
    if not catalog and g.get("variants"):
        catalog = [g["variants"]]
    add_numbered_lines(doc, catalog)
    add_field(doc, "Additional Info -", g["extra"])
    add_field(doc, "Examples on calendar (5 different dates) -", "")
    add_numbered_lines(doc, list(g.get("examples") or ()))
    doc.add_paragraph()
